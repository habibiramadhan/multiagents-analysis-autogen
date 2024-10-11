from .base_agent import BaseAgent
from typing import Dict, Any, List
import json
from datetime import datetime
from docx import Document
import pandas as pd

class ReporterAgent(BaseAgent):
    def __init__(self, name: str, llm_config: Dict[str, Any]):
        system_message = """Anda adalah seorang spesialis pelaporan. Tanggung jawab Anda mencakup:
        1. Membuat laporan analisis yang komprehensif
        2. Merangkum temuan utama
        3. Membuat laporan teknis dan non-teknis
        4. Menyertakan visualisasi yang relevan dalam laporan"""
        
        super().__init__(name=name, system_message=system_message, llm_config=llm_config)
        
    def create_report(
        self,
        analysis_results: Dict[str, Any],
        visualization_files: Dict[str, List[str]],
        report_type: str = "teknis"
    ) -> str:
        """Membuat laporan komprehensif dari hasil analisis."""
        try:
            doc = Document()
            
            # Menambahkan judul
            doc.add_heading(f'Laporan Analisis Data ({report_type.capitalize()})', 0)
            doc.add_paragraph(f'Digenerate pada: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # Menambahkan ringkasan eksekutif
            doc.add_heading('Ringkasan Eksekutif', level=1)
            summary = self._generate_summary(analysis_results)
            doc.add_paragraph(summary)
            
            # Menambahkan hasil analisis rinci
            if report_type == "teknis":
                self._add_technical_details(doc, analysis_results)
            else:
                self._add_business_insights(doc, analysis_results)
            
            # Menambahkan visualisasi
            doc.add_heading('Visualisasi', level=1)
            for viz_type, files in visualization_files.items():
                doc.add_heading(f'Visualisasi {viz_type.capitalize()}', level=2)
                for file in files:
                    doc.add_picture(file, width=6000000)  # ~6 inci
                    doc.add_paragraph(f'Gambar: {file.split("/")[-1]}')
            
            # Menyimpan laporan
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_path = f"output/reports/analysis_report_{report_type}_{timestamp}.docx"
            doc.save(report_path)
            
            return report_path
            
        except Exception as e:
            self.handle_error(e)
            return ""
            
    def _generate_summary(self, analysis_results: Dict[str, Any]) -> str:
        """Membuat ringkasan eksekutif dari hasil analisis."""        
        summary = []
        
        if "descriptive_statistics" in analysis_results:
            summary.append("Analisis Statistik: Metrik utama dihitung untuk semua variabel numerik.")
            
        if "correlation_analysis" in analysis_results:
            summary.append("Analisis Korelasi: Hubungan antar variabel diperiksa.")
            
        if "regression_analysis" in analysis_results:
            reg_results = analysis_results["regression_analysis"]
            r2 = reg_results.get("r_squared", 0)
            summary.append(f"Analisis Regresi: Model mencapai nilai R-kuadrat sebesar {r2:.2f}.")
            
        if "clustering_analysis" in analysis_results:
            cluster_results = analysis_results["clustering_analysis"]
            n_clusters = len(cluster_results.get("cluster_centers", {}))
            summary.append(f"Analisis Klastering: Data dibagi menjadi {n_clusters} kelompok yang berbeda.")
            
        return "\n\n".join(summary)
        
    def _add_technical_details(self, doc: Document, analysis_results: Dict[str, Any]):
        """Menambahkan rincian teknis ke laporan."""        
        for analysis_type, results in analysis_results.items():
            doc.add_heading(f'{analysis_type.replace("_", " ").title()}', level=1)
            
            if analysis_type == "descriptive_statistics":
                for var, stats in results.items():
                    doc.add_heading(f'Variabel: {var}', level=2)
                    for stat, value in stats.items():
                        doc.add_paragraph(f'{stat}: {value:.4f}')
                        
            elif analysis_type == "regression_analysis":
                doc.add_paragraph(f'R-kuadrat: {results["r_squared"]:.4f}')
                doc.add_heading('Koefisien', level=2)
                for feature, coef in results["coefficients"].items():
                    doc.add_paragraph(f'{feature}: {coef:.4f}')
                    
    def _add_business_insights(self, doc: Document, analysis_results: Dict[str, Any]):
        """Menambahkan wawasan bisnis ke laporan."""        
        doc.add_heading('Wawasan Utama', level=1)
        
        if "correlation_analysis" in analysis_results:
            doc.add_paragraph('Hubungan Utama:')
            correlations = pd.DataFrame(analysis_results["correlation_analysis"])
            strong_corrs = []
            
            for col1 in correlations.columns:
                for col2 in correlations.columns:
                    if col1 < col2:
                        corr = correlations.loc[col1, col2]
                        if abs(corr) > 0.7:
                            strong_corrs.append(
                                f"• Hubungan kuat {'positif' if corr > 0 else 'negatif'} "
                                f"antara {col1} dan {col2} (korelasi: {corr:.2f})"
                            )
                            
            for insight in strong_corrs:
                doc.add_paragraph(insight)
